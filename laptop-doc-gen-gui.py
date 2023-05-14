import sys
import os
from datetime import datetime
import re
import time
from docx.shared import Pt 
from PyQt6.QtWidgets import QApplication, QMainWindow,QComboBox, QMessageBox, QMenuBar, QFormLayout, QWidget, QLineEdit, QLabel, QPushButton, QSizePolicy
from PyQt6.QtGui import QAction, QIcon, QCursor, QFont
from PyQt6.QtCore import Qt, QSize
import docx # pip install python-docx
from deletedocs import DocFileDeleter


class LaptopDocGen(QMainWindow):
    def __init__(self):
        """
        __init__ an app that generates the Laptop Loan Welcome Letter and 2 copies of the Agreement document, with the student and device details inserted.

        The app will send the document to the default printer, and then ask user if they'd like to delete the .docx files. All document creations are logged to the log file to add one more point of reference when trying to resolve laptop loan logging-related issues.
        """
        super().__init__()
        self.setWindowTitle("Laptop Doc Gen v1.0.1")
        self.setWindowIcon(QIcon("img/laptop-icon.png"))
        self.setSizePolicy(QSizePolicy.Policy.Preferred, QSizePolicy.Policy.Preferred)

        self.thefont = "Roboto"
        self.logfile = "logs/log.txt"
        # Create a central widget with a form layout
        central_widget = QWidget(self)
        form_layout = QFormLayout(central_widget)
        form_layout.setContentsMargins(25, 35, 25, 50)
        form_layout.setSpacing(20)

        self.labels = {
            "name": QLabel("Student's name:", self),
            "serial": QLabel("Serial number:", self),
            "asset": QLabel("Asset tag:", self),
            "model": QLabel("Device model:", self),
        }
        self.line_edits = {
            "name": QLineEdit(self),
            # serial number,
            "serial": QLineEdit(self),
            # asset tag
            "asset": QLineEdit(self),
            # model of laptop (Dell Latitude 3310)
            "model": QLineEdit(self),
        }
        self.line_edits["name"].setPlaceholderText("Harry Potter (7xxxxxx)")
        self.line_edits["model"].setPlaceholderText("Dell Latitude 3310")

        for key in self.labels.keys():
            self.line_edits[key].setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Fixed)
            self.line_edits[key].setFixedHeight(40)
            
            form_layout.addRow(self.labels[key], self.line_edits[key])
        
        # create a 'send to printer' button
        submit_button = QPushButton("SEND TO PRINTER", self)
        submit_button.setFont(QFont("Roboto", 16, QFont.Weight.Bold.value))
        submit_button.setCursor(QCursor(Qt.CursorShape.PointingHandCursor))
        submit_button.clicked.connect(self.generate_docs)
        form_layout.addRow(submit_button)

        # Set the central widget
        self.setCentralWidget(central_widget)
    def generate_docs(self):
        """
        generate_docs gets the text from each of the textboxes, and inserts it into copies of the base documents.

        Creates the welcome and agreement documents by calling the create_welcome and create_agreement functions, then prints the documents by calling the print_docs function.
        """
        # get the values from the line edits
        name = self.line_edits["name"].text()
        serial = self.line_edits["serial"].text()
        asset = self.line_edits["asset"].text()
        model = self.line_edits["model"].text()
        for value in [name, serial, asset, model]:
            if value == "":
                return
        
        # path to the welcome / agreement base files
        self.welcome_file = './files/welcome.docx'
        self.agreement = './files/agreement.docx'
        try:
            self.welcome_doc = docx.Document(self.welcome_file)
            self.agree_doc = docx.Document(self.agreement)
        except Exception:
            popup = QMessageBox(self)
            popup.setText("Error: could not open the base files.")
            popup.setWindowTitle("Error")
            popup.setIcon(QMessageBox.Icon.Critical)                            
            popup.exec()
            return

        try:
            self.create_welcome(name)
            self.create_agreement(name, serial, asset, model)
            self.print_docs()

        except Exception:
            popup = QMessageBox(self)
            popup.button = QMessageBox.StandardButton.Ok
            popup.setText("Error: could not create the documents.")
            popup.setWindowTitle("Error")
            popup.setIcon(QMessageBox.Icon.Critical)                            
            if popup.exec() == QMessageBox.StandardButton.Ok:
                popup.close()
            return
        self.delete_docs_question()
    # user inputs name, then a word doc is created with students name inserted
    def create_welcome(self, student_name):
        # grab first sentence of the body/text of the base welcome doc
        first_sentence = self.welcome_doc.paragraphs[1]
        if re.match(r'\((\d{7})\)$', student_name):
        # The only variable text on the document
            welcome_string = f"Dear {student_name[:-12]}"
        else:
            welcome_string = f"Dear {student_name}"
        # rest of that chunk of text on the document
        rest_of_welcome = ": Welcome to Delaware Technical Community College!  This letter contains some helpful information pertaining to your laptop’s configuration of software and virtual support for assisting with technology issues."

        first_sentence.text = "" # erase the first sentence
        greeting = first_sentence.add_run(welcome_string) # add the new welcome string into the blank first sentence, and save the new 
        
        self.format_item(greeting,Pt(12), True, False) # sentence/python docx object to greeting variable so it can be used

        # in the formatting function
        # then add the rest of the paragraph, not bold
        rest = first_sentence.add_run(rest_of_welcome)
        self.format_item(rest, Pt(12), False, False)

        # save the welcome file
        self.welcome_filepath = f'{student_name}-welcome.docx'
        self.welcome_doc.save(self.welcome_filepath)
    def sizeHint(self):
        return QSize(500, 350)

    def create_agreement(self, student_name, servtag, assettag, dev):
        """
        create_agreement create the laptop loan agreement document that the student will sign upon receiving their laptop.

        _extended_summary_

        Args:
            student_name (str): student's name, followed by #700 inside parentheses
            servtag (str): serial number of the laptop/device
            assettag (str): asset tag of the laptop/device
            dev (str): make/model of the laptop/device
        """
        # deal with the first sentence - insert student's name
        first_sentence = self.agree_doc.paragraphs[2]
        start = "On the date of _____/_____/______, I, "

        first_sentence.text = ""
        started = first_sentence.add_run(start)
        self.format_item(started, Pt(12), False, False)
        # making the students name a bit bigger on the form, the different formatting is why there are 3 add_runs here instead of just one to add the whole first sentence (could be quicker way to do this)
        add_name = first_sentence.add_run(f'{student_name}')
        self.format_item(add_name, Pt(13), False, True)
        add_rest = first_sentence.add_run(', received the following computer equipment and accessories (“the Equipment”) from Delaware Technical Community College (“DTCC”):')
        self.format_item(add_rest, Pt(12), False, False)
        # inserting the device make/model, service tag and asset #
        two_underlines = self.agree_doc.paragraphs[3] # first have to target the chunk of text and erase it
        two_underlines.text = ""
        # add some blank space before the underlined make/model part starts
        two_underlines.add_run('         ')
        # check how many characters are in make/model string
        devlength = len(dev)
        # 56 spaces is the total number of spaces - underlines, make/model string included from start to end
        # subtracting the length of dev string input by user from 56, then dividing by 2 and rounding to whole number gives you the amount of underline you should have on either side of the device string to make it look even
        spaces = " " * round(((56 - devlength)/2))
        # add to document, spaces included
        make_model = two_underlines.add_run(f'{spaces}{dev}{spaces}')
        self.format_item(make_model, Pt(13), False, True)
        # put a tab of space in between the device description and the serial/asset
        two_underlines.add_run('\t')
        # format: <serial number> / <asset tag>, ex: J854MNB / 2003198
        assetstring = f"{servtag} / {assettag}"
        tags = two_underlines.add_run(f'{assetstring}')
        self.format_item(tags, Pt(13),False, True)

        # create new agreement file with student's name
        self.agreement_filepath = f'{student_name}-agreement.docx'
        self.agree_doc.save(self.agreement_filepath)
        
        dt = datetime.now()
        timestamp = dt.strftime("%B %d, %Y - %I:%M %p")
        
        with open(self.logfile, 'a+') as thelog:
            thelog.write(f"{timestamp} : {student_name}  {assetstring} ({dev})")
        
        
    def print_docs(self):
        """
        print_docs sends the documents to default printer using the python os library's 'startfile' function with the 'print' argument
        """
        os.startfile(self.welcome_filepath, "print")
        for _ in range(2):
            os.startfile(self.agreement_filepath, "print")
            time.sleep(1)
    def format_item(self, item, item_size, bold, underline):
        """
        format_item used to efficiently format certain text items in the documents

        Args:
            item (_type_): _description_
            item_size (_type_): _description_
            bold (_type_): _description_
            underline (_type_): _description_
        """
        item.font.name = self.thefont
        item.font.size = item_size
        item.bold = bold
        item.underline = underline
    def delete_docs_question(self):
        """
        delete_docs_question ask the user if they want to delete the .docx files after printing, it's a good idea to do this because the files have student's 700 numbers in them.
        """
        # create a message box to ask user if they want to delete the .docx files
        msg = QMessageBox()
        msg.setWindowTitle("Delete .docx files?")
        msg.setText("Documents sent to printer.\nWould you like to delete the .docx files and clear form?")
        msg.setIcon(QMessageBox.Icon.Question)
        msg.setStandardButtons(QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        msg.setDefaultButton(QMessageBox.StandardButton.Yes)
        if msg.exec() == QMessageBox.StandardButton.Yes:
            deleter = DocFileDeleter()
            deleter.delete_doc_files()
        for key in self.line_edits.keys():
            self.line_edits[key].setText("")

    def write_log(self, log_string):
        """
        write_log writes supplied string of text to the log file, with current date/time

        Args:
            log_string (str): string of text to be written to log file
        """
        # get date time string
        now = datetime.now()
        dt_string = now.strftime("%m/%d/%Y %H:%M:%S")
        # write to log file
        with open('log.txt', 'a') as f:
            self.logfile.write(f'{dt_string} - {log_string}')
if __name__ == '__main__':
    app = QApplication(sys.argv)
    # set the theme of app, right now there's just the dark theme since it's a small app
    with open('themes/dark.qss', 'r') as darkfile:
        darktheme = darkfile.read()
        
        app.setStyleSheet(darktheme)
    window = LaptopDocGen()
    window.show()
    sys.exit(app.exec())
